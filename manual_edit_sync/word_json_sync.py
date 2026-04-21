from __future__ import annotations

import json
import time
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from docx import Document


@dataclass
class SyncResult:
    report_path: str
    srs_json_path: str
    tests_json_path: str
    srs_count: int
    tests_count: int
    srs_updated: int
    tests_updated: int
    timestamp: float


def _parse_desc_and_objective(raw_text: str) -> Tuple[str, str]:
    lines = [line.strip() for line in raw_text.splitlines() if line.strip()]
    if not lines:
        return "", ""

    desc = ""
    obj = ""
    for line in lines:
        normalized = line.replace("：", ":")
        lower = normalized.lower()
        if lower.startswith("desc:") or lower.startswith("description:") or normalized.startswith("描述:"):
            desc = normalized.split(":", 1)[1].strip()
        elif lower.startswith("obj:") or lower.startswith("objective:") or normalized.startswith("目标:"):
            obj = normalized.split(":", 1)[1].strip()

    if not desc and not obj:
        # If users removed prefixes manually, keep full text as description.
        return raw_text.strip(), ""

    return desc, obj


def _parse_test_data(raw_text: str) -> Any:
    text = raw_text.strip()
    if not text:
        return ""

    lines = [line.strip().replace("：", ":") for line in text.splitlines() if line.strip()]
    if not lines:
        return ""

    as_dict: Dict[str, str] = {}
    for line in lines:
        if ":" not in line:
            return text
        key, value = line.split(":", 1)
        key = key.strip()
        value = value.strip()
        if not key:
            return text
        as_dict[key] = value

    return as_dict


def _load_json(path: Path) -> List[Dict[str, Any]]:
    if not path.exists():
        return []
    with path.open("r", encoding="utf-8") as f:
        data = json.load(f)
        return data if isinstance(data, list) else []


def _dump_json(path: Path, data: List[Dict[str, Any]]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


def _resolve_default_json_paths(report_path: Path) -> Tuple[Path, Path]:
    stem = report_path.stem
    if stem.endswith("_功能测试报告"):
        base = stem[: -len("_功能测试报告")]
    else:
        # Fallback for non-standard names.
        base = stem

    srs_path = report_path.with_name(f"{base}_SRS.json")
    tests_path = report_path.with_name(f"{base}_TestRequirements.json")
    return srs_path, tests_path


def _parse_tables_from_report(report_path: Path) -> Tuple[List[Dict[str, str]], List[Dict[str, Any]]]:
    doc = Document(str(report_path))
    if len(doc.tables) < 2:
        raise ValueError("Word report format mismatch: expected at least two tables (SRS and Tests).")

    srs_table = doc.tables[0]
    tests_table = doc.tables[1]

    srs_rows: List[Dict[str, str]] = []
    for row in srs_table.rows[1:]:
        cells = [cell.text.strip() for cell in row.cells]
        if len(cells) < 4:
            continue
        if not any(cells):
            continue
        srs_rows.append(
            {
                "id": cells[0],
                "name": cells[1],
                "requirement": cells[2],
                "type": cells[3],
            }
        )

    test_rows: List[Dict[str, Any]] = []
    for row in tests_table.rows[1:]:
        cells = [cell.text.strip() for cell in row.cells]
        if len(cells) < 5:
            continue
        if not any(cells):
            continue

        description, objective = _parse_desc_and_objective(cells[2])
        test_rows.append(
            {
                "id": cells[0],
                "name": cells[1],
                "description": description,
                "objective": objective,
                "test_data": _parse_test_data(cells[3]),
                "expected_result": cells[4],
            }
        )

    return srs_rows, test_rows


def _merge_by_id(
    edited_rows: List[Dict[str, Any]],
    original_rows: List[Dict[str, Any]],
    fields_to_update: List[str],
) -> Tuple[List[Dict[str, Any]], int]:
    original_by_id = {str(item.get("id", "")).strip(): item for item in original_rows}

    merged: List[Dict[str, Any]] = []
    updated = 0
    for row in edited_rows:
        row_id = str(row.get("id", "")).strip()
        source = dict(original_by_id.get(row_id, {}))

        for key in fields_to_update:
            source[key] = row.get(key, "")

        if source != original_by_id.get(row_id, {}):
            updated += 1
        merged.append(source)

    return merged, updated


def sync_word_to_json(
    report_path: str,
    srs_json_path: Optional[str] = None,
    tests_json_path: Optional[str] = None,
) -> SyncResult:
    report = Path(report_path)
    if not report.exists():
        raise FileNotFoundError(f"Report file not found: {report}")

    if srs_json_path and tests_json_path:
        srs_path = Path(srs_json_path)
        tests_path = Path(tests_json_path)
    else:
        srs_path, tests_path = _resolve_default_json_paths(report)

    parsed_srs, parsed_tests = _parse_tables_from_report(report)

    old_srs = _load_json(srs_path)
    old_tests = _load_json(tests_path)

    new_srs, srs_updated = _merge_by_id(
        edited_rows=parsed_srs,
        original_rows=old_srs,
        fields_to_update=["id", "name", "requirement", "type"],
    )
    new_tests, tests_updated = _merge_by_id(
        edited_rows=parsed_tests,
        original_rows=old_tests,
        fields_to_update=["id", "name", "description", "objective", "test_data", "expected_result"],
    )

    _dump_json(srs_path, new_srs)
    _dump_json(tests_path, new_tests)

    return SyncResult(
        report_path=str(report),
        srs_json_path=str(srs_path),
        tests_json_path=str(tests_path),
        srs_count=len(new_srs),
        tests_count=len(new_tests),
        srs_updated=srs_updated,
        tests_updated=tests_updated,
        timestamp=time.time(),
    )
