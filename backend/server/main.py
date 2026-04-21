import sys
import os
import time
import shutil
import uuid
import asyncio
from dataclasses import asdict
from pathlib import Path
from typing import Any, Dict, Optional, List
from fastapi import FastAPI, File, UploadFile, Form, BackgroundTasks, HTTPException
from fastapi.staticfiles import StaticFiles
from fastapi.responses import JSONResponse, FileResponse
from pydantic import BaseModel
from docx import Document

# Add project root to path
ROOT_DIR = Path(__file__).resolve().parent.parent.parent
if str(ROOT_DIR) not in sys.path:
    sys.path.insert(0, str(ROOT_DIR))

from Req.demo.run_multi_model_unified6_demo import process_existing_app_combo, AVAILABLE_COMBOS
from Req.tools.parse_flow import preprocess_apk
from Req.config.RunConfig import OUTPUT_ROOT
from Req.tools.report_generator import generate_report
from Req.tools.zip_utils import create_zip
from manual_edit_sync.word_json_sync import sync_word_to_json

app = FastAPI()

# Storage configuration
STORAGE_DIR = ROOT_DIR / "storage"
UPLOAD_DIR = STORAGE_DIR / "uploads"
DOWNLOAD_DIR = STORAGE_DIR / "downloads"

UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
DOWNLOAD_DIR.mkdir(parents=True, exist_ok=True)

import json

# Global Configuration
GLOBAL_API_KEY = os.environ.get("DASHSCOPE_API_KEY", "")

# Job Store
JOBS_FILE = ROOT_DIR / "jobs.json"
jobs = {}
manual_edit_watch_tasks: Dict[str, asyncio.Task] = {}

def load_jobs():
    global jobs
    if JOBS_FILE.exists():
        try:
            with open(JOBS_FILE, "r", encoding="utf-8") as f:
                jobs = json.load(f)
            print(f"Loaded {len(jobs)} jobs from {JOBS_FILE}")
        except Exception as e:
            print(f"Failed to load jobs: {e}")
            jobs = {}

def save_jobs():
    try:
        with open(JOBS_FILE, "w", encoding="utf-8") as f:
            json.dump(jobs, f, indent=2, ensure_ascii=False)
    except Exception as e:
        print(f"Failed to save jobs: {e}")

# Load jobs on startup
load_jobs()

class JobStatus:
    PENDING = "pending"
    RUNNING = "running"
    COMPLETED = "completed"
    FAILED = "failed"

class GenerateRequest(BaseModel):
    job_id: str
    combo: str
    lang: str = "zh"
    api_key: Optional[str] = None
    repo_link: Optional[str] = None
    app_intro: Optional[str] = None
    app_name: Optional[str] = None

class ConfigRequest(BaseModel):
    api_key: str

class ManualEditSyncRequest(BaseModel):
    job_id: Optional[str] = None
    report_path: Optional[str] = None
    srs_json_path: Optional[str] = None
    tests_json_path: Optional[str] = None

class ManualEditWatchRequest(ManualEditSyncRequest):
    poll_interval_seconds: float = 2.0


class ManualEditSaveRequest(BaseModel):
    job_id: str
    srs_rows: List[Dict[str, str]]
    test_rows: List[Dict[str, Any]]

@app.get("/api/config")
def get_config():
    return {"has_api_key": bool(GLOBAL_API_KEY)}

def _resolve_report_path(req: ManualEditSyncRequest) -> str:
    if req.report_path:
        return req.report_path

    if req.job_id:
        job = jobs.get(req.job_id)
        if not job:
            raise HTTPException(status_code=404, detail="Job not found")
        if job.get("status") != JobStatus.COMPLETED:
            raise HTTPException(status_code=400, detail="Job not completed")
        report_path = job.get("result", {}).get("report")
        if not report_path:
            raise HTTPException(status_code=400, detail="Report path not found in job result")
        return report_path

    raise HTTPException(status_code=400, detail="Either job_id or report_path is required")


def _test_desc_objective_to_cell(description: str, objective: str) -> str:
    desc = (description or "").strip()
    obj = (objective or "").strip()
    if desc and obj:
        return f"Desc: {desc}\nObj: {obj}"
    if desc:
        return f"Desc: {desc}"
    if obj:
        return f"Obj: {obj}"
    return ""


def _test_data_to_cell(test_data: Any) -> str:
    if isinstance(test_data, dict):
        return "\n".join([f"{k}: {v}" for k, v in test_data.items()])
    if test_data is None:
        return ""
    return str(test_data)


def _parse_desc_objective_from_cell(raw_text: str) -> Dict[str, str]:
    lines = [line.strip() for line in (raw_text or "").splitlines() if line.strip()]
    if not lines:
        return {"description": "", "objective": ""}

    description = ""
    objective = ""
    for line in lines:
        normalized = line.replace("：", ":")
        lower = normalized.lower()
        if lower.startswith("desc:") or lower.startswith("description:") or normalized.startswith("描述:"):
            description = normalized.split(":", 1)[1].strip()
        elif lower.startswith("obj:") or lower.startswith("objective:") or normalized.startswith("目标:"):
            objective = normalized.split(":", 1)[1].strip()

    if not description and not objective:
        description = (raw_text or "").strip()

    return {"description": description, "objective": objective}


def _parse_report_for_editor(report_path: str) -> Dict[str, Any]:
    doc = Document(report_path)
    if len(doc.tables) < 2:
        raise ValueError("Word report format mismatch: expected at least two tables (SRS and Tests).")

    srs_table = doc.tables[0]
    tests_table = doc.tables[1]

    srs_rows: List[Dict[str, str]] = []
    for row in srs_table.rows[1:]:
        cells = [cell.text.strip() for cell in row.cells]
        if len(cells) < 4:
            continue
        if not any(cells):
            continue
        srs_rows.append(
            {
                "id": cells[0],
                "name": cells[1],
                "requirement": cells[2],
                "type": cells[3],
            }
        )

    test_rows: List[Dict[str, Any]] = []
    for row in tests_table.rows[1:]:
        cells = [cell.text.strip() for cell in row.cells]
        if len(cells) < 5:
            continue
        if not any(cells):
            continue
        parsed = _parse_desc_objective_from_cell(cells[2])
        test_rows.append(
            {
                "id": cells[0],
                "name": cells[1],
                "description": parsed["description"],
                "objective": parsed["objective"],
                "test_data": cells[3],
                "expected_result": cells[4],
            }
        )

    return {
        "report_path": report_path,
        "report_name": Path(report_path).name,
        "srs_rows": srs_rows,
        "test_rows": test_rows,
    }


def _replace_report_tables(report_path: str, srs_rows: List[Dict[str, str]], test_rows: List[Dict[str, Any]]) -> None:
    doc = Document(report_path)
    if len(doc.tables) < 2:
        raise ValueError("Word report format mismatch: expected at least two tables (SRS and Tests).")

    srs_table = doc.tables[0]
    tests_table = doc.tables[1]

    srs_tbl = srs_table._tbl
    for row in list(srs_table.rows[1:]):
        srs_tbl.remove(row._tr)

    for item in srs_rows:
        row = srs_table.add_row().cells
        row[0].text = str(item.get("id", ""))
        row[1].text = str(item.get("name", ""))
        row[2].text = str(item.get("requirement", ""))
        row[3].text = str(item.get("type", ""))

    tests_tbl = tests_table._tbl
    for row in list(tests_table.rows[1:]):
        tests_tbl.remove(row._tr)

    for item in test_rows:
        row = tests_table.add_row().cells
        row[0].text = str(item.get("id", ""))
        row[1].text = str(item.get("name", ""))
        row[2].text = _test_desc_objective_to_cell(
            description=str(item.get("description", "")),
            objective=str(item.get("objective", "")),
        )
        row[3].text = _test_data_to_cell(item.get("test_data"))
        row[4].text = str(item.get("expected_result", ""))

    doc.save(report_path)


@app.get("/api/manual-edit/editor/{job_id}")
def manual_edit_editor_load(job_id: str):
    req = ManualEditSyncRequest(job_id=job_id)
    report_path = _resolve_report_path(req)
    try:
        payload = _parse_report_for_editor(report_path)
        payload["ok"] = True
        payload["job_id"] = job_id
        return payload
    except Exception as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc


@app.post("/api/manual-edit/editor/save")
def manual_edit_editor_save(req: ManualEditSaveRequest):
    try:
        report_path = _resolve_report_path(ManualEditSyncRequest(job_id=req.job_id))
        _replace_report_tables(report_path, req.srs_rows, req.test_rows)
        sync_result = sync_word_to_json(report_path=report_path)
        payload = asdict(sync_result)
        payload["ok"] = True
        payload["job_id"] = req.job_id
        payload["message"] = "Word report updated and JSON synced"
        return payload
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc


@app.post("/api/manual-edit/sync")
def manual_edit_sync(req: ManualEditSyncRequest):
    try:
        report_path = _resolve_report_path(req)
        result = sync_word_to_json(
            report_path=report_path,
            srs_json_path=req.srs_json_path,
            tests_json_path=req.tests_json_path,
        )
        payload = asdict(result)
        payload["ok"] = True
        return payload
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc


async def _manual_edit_watch_loop(req: ManualEditWatchRequest) -> None:
    report_path = _resolve_report_path(req)
    last_mtime = os.path.getmtime(report_path)

    while True:
        await asyncio.sleep(req.poll_interval_seconds)
        try:
            new_mtime = os.path.getmtime(report_path)
            if new_mtime > last_mtime:
                sync_word_to_json(
                    report_path=report_path,
                    srs_json_path=req.srs_json_path,
                    tests_json_path=req.tests_json_path,
                )
                last_mtime = new_mtime
        except FileNotFoundError:
            continue


@app.post("/api/manual-edit/watch/start")
async def manual_edit_watch_start(req: ManualEditWatchRequest):
    if req.poll_interval_seconds < 0.5 or req.poll_interval_seconds > 30:
        raise HTTPException(status_code=400, detail="poll_interval_seconds must be between 0.5 and 30")

    report_path = _resolve_report_path(req)

    task = manual_edit_watch_tasks.get(report_path)
    if task and not task.done():
        return {"ok": True, "message": "Watcher already running", "report_path": report_path}

    task = asyncio.create_task(_manual_edit_watch_loop(req))
    manual_edit_watch_tasks[report_path] = task
    return {"ok": True, "message": "Watcher started", "report_path": report_path}


@app.post("/api/manual-edit/watch/stop")
async def manual_edit_watch_stop(req: ManualEditSyncRequest):
    report_path = _resolve_report_path(req)
    task = manual_edit_watch_tasks.get(report_path)
    if not task:
        return {"ok": True, "message": "Watcher not found", "report_path": report_path}

    task.cancel()
    return {"ok": True, "message": "Watcher stopped", "report_path": report_path}

@app.post("/api/config")
def set_config(config: ConfigRequest):
    global GLOBAL_API_KEY
    GLOBAL_API_KEY = config.api_key
    # Also set in env for immediate use if needed by other modules
    os.environ["DASHSCOPE_API_KEY"] = GLOBAL_API_KEY
    return {"message": "Configuration updated"}

@app.get("/api/combos")
def get_combos():
    return {"combos": AVAILABLE_COMBOS}

@app.post("/api/upload")
async def upload_file(file: UploadFile = File(...), app_name: Optional[str] = Form(None)):
    job_id = str(uuid.uuid4())
    job_dir = UPLOAD_DIR / job_id
    job_dir.mkdir(parents=True, exist_ok=True)
    
    file_path = job_dir / "app.apk"
    with open(file_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
    
    # Create meta info
    jobs[job_id] = {
        "job_id": job_id,
        "status": JobStatus.PENDING,
        "created_at": str(time.time()),
        "file_path": str(file_path),
        "original_filename": file.filename,
        "app_name": app_name or "Unknown",
        "combo": "",
        "lang": ""
    }
    save_jobs()
    
    return {
        "job_id": job_id,
        "message": "File uploaded successfully",
        "app_name": jobs[job_id]["app_name"],
        "original_filename": jobs[job_id]["original_filename"]
    }

def process_job_task(job_id: str, combo: str, lang: str, repo_link: str, app_intro: str, api_key: str):
    try:
        jobs[job_id]["status"] = JobStatus.RUNNING
        jobs[job_id]["combo"] = combo
        jobs[job_id]["lang"] = lang
        save_jobs()
        
        job_info = jobs[job_id]
        file_path = Path(job_info["file_path"])
        
        # Set API Key priority: Request > Global
        effective_api_key = api_key if api_key else GLOBAL_API_KEY
        
        if effective_api_key:
            os.environ["DASHSCOPE_API_KEY"] = effective_api_key
        
        # Preprocess APK
        work_dir = file_path.parent / "work"
        work_dir.mkdir(exist_ok=True)
        
        pp = preprocess_apk(str(file_path), str(work_dir))
        
        if not pp.get("ok"):
            raise Exception(f"APK Preprocessing failed: {pp.get('message')}")
            
        app_dir = Path(pp.get("app_dir"))
        
        # Run Generation
        result = process_existing_app_combo(
            app_dir=app_dir,
            combo_label=combo,
            lang=lang
        )
        
        if not result.get("ok"):
             raise Exception(f"Generation failed: {result.get('message')}")
             
        # Generate Report and Zip
        srs_path = result.get("srs")
        tests_path = result.get("tests")
        test_json_path = result.get("test_json")
        app_name = result.get("app")
        
        # jobs[job_id]["app_name"] = app_name # Keep original filename as app_name
        
        # Determine output directory (same as generated files)
        combo_dir = Path(srs_path).parent
        
        # 1. Generate Word Report
        report_filename = f"{app_name}_功能测试报告.docx"
        report_path = combo_dir / report_filename
        generate_report(srs_path, tests_path, str(report_path), lang=lang)
        
        # 2. Generate Zip
        zip_filename = f"{app_name}_data.zip"
        zip_path = combo_dir / zip_filename
        
        files_to_zip = [srs_path, tests_path]
        if test_json_path and os.path.exists(test_json_path):
            files_to_zip.append(test_json_path)
            
        create_zip(files_to_zip, str(zip_path))
        
        # Update result with new files
        result["report"] = str(report_path)
        result["zip"] = str(zip_path)

        # Result paths are absolute.
        jobs[job_id]["result"] = result
        jobs[job_id]["status"] = JobStatus.COMPLETED
        save_jobs()
        
    except Exception as e:
        jobs[job_id]["status"] = JobStatus.FAILED
        jobs[job_id]["error"] = str(e)
        save_jobs()
        print(f"Job {job_id} failed: {e}")

@app.post("/api/generate")
async def start_generation(req: GenerateRequest, background_tasks: BackgroundTasks):
    job_id = req.job_id
    if job_id not in jobs:
        raise HTTPException(status_code=404, detail="Job not found")
    
    # Update app_name if provided (in case user entered it after upload)
    if req.app_name:
        jobs[job_id]["app_name"] = req.app_name
        save_jobs()
        
    background_tasks.add_task(
        process_job_task, 
        job_id, 
        req.combo, 
        req.lang, 
        req.repo_link, 
        req.app_intro, 
        req.api_key
    )
    
    return {"job_id": job_id, "status": "queued"}

@app.get("/api/status/{job_id}")
async def get_status(job_id: str):
    if job_id not in jobs:
        raise HTTPException(status_code=404, detail="Job not found")
    return jobs[job_id]

@app.get("/api/history")
async def get_history():
    # Return list of jobs sorted by creation time (newest first)
    # We use 'created_at' which is a float string timestamp
    
    job_list = []
    for jid, job in jobs.items():
        # Create a summary object
        summary = {
            "job_id": jid,
            "status": job.get("status"),
            "created_at": job.get("created_at"),
            "original_filename": job.get("original_filename"),
            "app_name": job.get("app_name", "Unknown"),
            "combo": job.get("combo"),
            "lang": job.get("lang"),
            "error": job.get("error")
        }
        # Include result paths if completed
        if job.get("status") == JobStatus.COMPLETED:
            summary["result"] = job.get("result")
            
        job_list.append(summary)
        
    # Sort by created_at desc
    job_list.sort(key=lambda x: float(x.get("created_at", 0)), reverse=True)
    
    return {"history": job_list}

@app.get("/api/download/{job_id}/{file_type}")
async def download_result(job_id: str, file_type: str):
    if job_id not in jobs:
        raise HTTPException(status_code=404, detail="Job not found")
    
    job = jobs[job_id]
    if job["status"] != JobStatus.COMPLETED:
        raise HTTPException(status_code=400, detail="Job not completed")
        
    result = job.get("result", {})
    file_path = None
    
    if file_type == "srs":
        file_path = result.get("srs")
    elif file_type == "tests":
        file_path = result.get("tests")
    elif file_type == "test_json":
        file_path = result.get("test_json")
    elif file_type == "report":
        file_path = result.get("report")
    elif file_type == "zip":
        file_path = result.get("zip")
        
    if not file_path or not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="File not found")
        
    return FileResponse(file_path, filename=Path(file_path).name)

# Serve Static Files (Mount last to avoid blocking APIs)
app.mount("/", StaticFiles(directory=ROOT_DIR / "frontend", html=True), name="static")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8001)
