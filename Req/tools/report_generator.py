import json
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import RGBColor

def generate_report(srs_path: str, tests_path: str, output_path: str, lang: str = 'zh'):
    """
    Generates a Word report from SRS and Test Requirements JSON files.
    """
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # Read JSON data
    srs_data = []
    tests_data = []
    
    try:
        with open(srs_path, 'r', encoding='utf-8') as f:
            srs_data = json.load(f)
    except Exception as e:
        print(f"Error reading SRS file: {e}")
        
    try:
        with open(tests_path, 'r', encoding='utf-8') as f:
            tests_data = json.load(f)
    except Exception as e:
        print(f"Error reading Tests file: {e}")

    # Title
    title_text = "功能测试报告" if lang == 'zh' else "Function Test Report"
    heading = doc.add_heading(title_text, 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 1. Software Requirements
    srs_title = "1. 软件需求 (Software Requirements)"
    doc.add_heading(srs_title, level=1)
    
    if srs_data:
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'ID'
        hdr_cells[1].text = 'Name'
        hdr_cells[2].text = 'Requirement'
        hdr_cells[3].text = 'Type'
        
        # Make header bold and shaded (optional, keep simple for now)
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

        for item in srs_data:
            row_cells = table.add_row().cells
            row_cells[0].text = str(item.get('id', ''))
            row_cells[1].text = str(item.get('name', ''))
            row_cells[2].text = str(item.get('requirement', ''))
            row_cells[3].text = str(item.get('type', ''))
    else:
        doc.add_paragraph("No software requirements data available.")

    doc.add_page_break()

    # 2. Test Requirements
    tests_title = "2. 测试需求 (Test Requirements)"
    doc.add_heading(tests_title, level=1)
    
    if tests_data:
        # For tests, a table might be too wide. Let's use paragraphs or a vertical table per test.
        # But the user asked for format reference. Assuming list style or table.
        # Let's try a table but maybe fewer columns or just a list of tables.
        # A big table is usually standard.
        
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'ID'
        hdr_cells[1].text = 'Name'
        hdr_cells[2].text = 'Description/Objective'
        hdr_cells[3].text = 'Test Data'
        hdr_cells[4].text = 'Expected Result'
        
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    
        for item in tests_data:
            row_cells = table.add_row().cells
            row_cells[0].text = str(item.get('id', ''))
            row_cells[1].text = str(item.get('name', ''))
            
            # Combine Desc and Obj
            desc = item.get('description', '')
            obj = item.get('objective', '')
            content = f"Desc: {desc}\nObj: {obj}" if obj else desc
            row_cells[2].text = content
            
            # Test Data
            td = item.get('test_data', '')
            if isinstance(td, dict):
                td_str = "\n".join([f"{k}: {v}" for k, v in td.items()])
            else:
                td_str = str(td)
            row_cells[3].text = td_str
            
            # Expected Result
            row_cells[4].text = str(item.get('expected_result', ''))
            
    else:
        doc.add_paragraph("No test requirements data available.")
        
    doc.save(output_path)
    return output_path

if __name__ == "__main__":
    # Test
    pass
